
############### START STUFF #########################################################################
#### if we want to remove the database and start over. Esp if i add more documents. and we need to split and embed again
# rm -rf /home/rchak007/github/lausdLLM/chromaDocs/chroma/
#######################################################################################################


from docx import Document
from langchain.schema import Document as LC_Document
from docx.oxml import CT_Tbl, CT_P

from langchain.vectorstores import Chroma
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI

import re
from langchain.schema import Document as LC_Document

import os
import openai
import sys
sys.path.append('../..')
from dotenv import load_dotenv, find_dotenv
from docx import Document

# Load environment variables from .env file
load_dotenv(find_dotenv())

openai.api_key  = os.environ['OPENAI_API_KEY']



def extract_text_from_word(file_path):
    """
    Reads a Word document and extracts text, bullet points, and tables, preserving their original order.

    Args:
        file_path (str): Path to the Word document.

    Returns:
        str: Extracted and formatted text, including bullet points and tables.
    """
    doc = Document(file_path)
    extracted_text = []

    # Iterate through elements in the document in order (paragraphs and tables)
    for element in doc.element.body:
        if isinstance(element, CT_P):  # It's a paragraph
            para = element.text.strip()
            if para:
                extracted_text.append(para)
        elif isinstance(element, CT_Tbl):  # It's a table
            table = next(t for t in doc.tables if t._element is element)
            table_text = []

            # Add table header marker
            # table_text.append(f"\n=== Table {len(extracted_text) + 1} ===\n")

            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                if any(row_text):  # Avoid appending empty rows
                    table_text.append(" | ".join(row_text))  # Format row with separators

            if table_text:
                extracted_text.append("\n".join(table_text))

    return "\n\n".join(extracted_text)


# def extract_text_from_word(file_path):
#     doc = Document(file_path)
    
#     # Check table count first
#     print(f"Total tables found: {len(doc.tables)}")

#     for table_idx, table in enumerate(doc.tables):
#         print(f"\n=== Table {table_idx + 1} ===")  # Show table number
        
#         for row in table.rows[:5]:  # Print first 5 rows
#             print([cell.text.strip() for cell in row.cells])





def read_extract_Langdoc(word_file_path):
    """
    Reads a Word file, extracts structured text, and converts it to a LangChain Document.

    Args:
        word_file_path (str): Path to the Word document.

    Returns:
        list: A list containing a single LangChain Document object.
    """
    print(f"Processing document: {word_file_path}")
    
    # Extract structured text from Word
    extracted_text = extract_text_from_word(word_file_path)
    print("Extracted Text length: ", len(extracted_text))

    # Tested below that tables are now good.
    # print("Extracted Text Preview:", extracted_text[:15745], "\n")  # Preview first 500 chars

    print("Extracted Text Preview:", extracted_text[:500], "\n")  # Preview first 500 chars

    # Convert extracted text into LangChain document format
    word_docs = [LC_Document(page_content=extracted_text, metadata={"source": word_file_path})]

    print(f"Total Documents Loaded: {len(word_docs)}")
    print("LangChain Document Preview:", word_docs[0].page_content[:500], "\n")

    return word_docs

def load_db(word_docs, k,  chain_type="stuff", remove=False):
    """
    Processes extracted Word document text, chunks it, and loads it into Chroma vector DB.
    
    Args:
        word_docs (list): Extracted documents from Word.
        chain_type (str): Type of retrieval chain.
        k (int): Number of retrieval results.
        remove (bool): If True, deletes the existing Chroma DB before recreating it.

    Returns:
        ConversationalRetrievalChain: A retriever-powered chatbot chain.
    """
    persist_directory = "/home/rchak007/github/lausdLLM/chromaDocs/chroma/"

    # If remove=True, delete the existing Chroma database
    if remove and os.path.exists(persist_directory):
        print("🛑 Removing existing Chroma DB...")
        for root, dirs, files in os.walk(persist_directory, topdown=False):
            for file in files:
                os.remove(os.path.join(root, file))
            for dir in dirs:
                os.rmdir(os.path.join(root, dir))
        os.rmdir(persist_directory)
        print("✅ Chroma DB removed.")

    # Check if Chroma DB exists
    if os.path.exists(persist_directory) and os.listdir(persist_directory):
        print("🔄 Loading existing Chroma vector database...")
        db = Chroma(persist_directory=persist_directory, embedding_function=OpenAIEmbeddings())
    else:
        print("⚡ Creating new Chroma vector database...")

        # Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        docs = text_splitter.split_documents(word_docs)  # ✅ Use `word_docs` directly

        # Create vector store
        embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
        db = Chroma.from_documents(docs, embeddings, persist_directory=persist_directory)
        db.persist()  # ✅ Save for future use

    # Create retriever and chat chain
    retriever = db.as_retriever(search_type="mmr", search_kwargs={"k": k, "fetch_k": 20})
    qa = ConversationalRetrievalChain.from_llm(
        llm=ChatOpenAI(model_name="gpt-4-turbo", temperature=0), 
        chain_type=chain_type, 
        retriever=retriever, 
        return_source_documents=True,
        return_generated_question=True,
    )
    
    return qa


import os
from langchain.vectorstores import Chroma
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI

def chunk_documents(word_docs, chunk_size=1000, chunk_overlap=150):
    """
    Splits extracted text into chunks for embedding.

    Args:
        word_docs (list): Extracted documents from Word.
        chunk_size (int): Size of each chunk.
        chunk_overlap (int): Overlapping characters between chunks.

    Returns:
        list: Chunked documents.
    """
    print("📏 Splitting documents into chunks...")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return text_splitter.split_documents(word_docs)


def chunk_TokenTextSplitter(char_chunks, token_chunk_size=256, token_chunk_overlap=50):
    """
    Further refines character-based chunks into token-based chunks using TokenTextSplitter.

    Args:
        char_chunks (list): Character-based chunked documents.
        token_chunk_size (int): Size of each chunk (in tokens).
        token_chunk_overlap (int): Overlapping tokens between chunks.

    Returns:
        list: Token-based chunked documents.
    """
    print("📏 Refining chunks using TokenTextSplitter...")
    token_splitter = TokenTextSplitter(
        chunk_size=token_chunk_size, 
        chunk_overlap=token_chunk_overlap
    )
    token_chunks = token_splitter.split_documents(char_chunks)
    print(f"🔍 Total Token-Based Chunks: {len(token_chunks)}")
    
    return token_chunks


def create_vector_store(docs, persist_directory, remove=False):
    """
    Embeds and stores chunks into ChromaDB.

    Args:
        docs (list): Chunked documents.
        persist_directory (str): Path to store ChromaDB.
        remove (bool): If True, removes existing ChromaDB before recreating.

    Returns:
        Chroma: Vector store database.
    """
    if remove and os.path.exists(persist_directory):
        print("🛑 Removing existing Chroma DB...")
        for root, dirs, files in os.walk(persist_directory, topdown=False):
            for file in files:
                os.remove(os.path.join(root, file))
            for dir in dirs:
                os.rmdir(os.path.join(root, dir))
        os.rmdir(persist_directory)
        print("✅ Chroma DB removed.")

    if os.path.exists(persist_directory) and os.listdir(persist_directory):
        print("🔄 Loading existing Chroma vector database...")
        return Chroma(persist_directory=persist_directory, embedding_function=OpenAIEmbeddings())
    
    print("⚡ Creating new Chroma vector database...")
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
    db = Chroma.from_documents(docs, embeddings, persist_directory=persist_directory)
    db.persist()
    return db


def initialize_retriever_and_qa(db, chain_type="stuff", k=10):
    """
    Creates the retriever and initializes the chatbot.

    Args:
        db (Chroma): Vector store database.
        chain_type (str): Type of retrieval chain.
        k (int): Number of retrieval results.

    Returns:
        ConversationalRetrievalChain: Chatbot retrieval chain.
    """
    retriever = db.as_retriever(search_type="mmr", search_kwargs={"k": k, "fetch_k": 20})
    return ConversationalRetrievalChain.from_llm(
        llm=ChatOpenAI(model_name="gpt-4-turbo", temperature=0),
        chain_type=chain_type,
        retriever=retriever,
        return_source_documents=True,
        return_generated_question=True,
    )


def examine_chunks(chunks, num_preview=5):
    """
    Examines the chunked documents by:
    - Counting total chunks
    - Previewing a few chunks
    - Identifying the largest & smallest chunks

    Args:
        chunks (list): List of chunked documents.
        num_preview (int): Number of chunks to preview (default is 5).
    """
    print(f"\n🔢 Total Chunks Created: {len(chunks)}\n")

    # Show first few chunks
    for i, chunk in enumerate(chunks[:num_preview]):  
        # print(f"\n📝 Chunk {i+1} (Length: {len(chunk.page_content)} characters):\n{chunk.page_content[:1500]}...\n")
        print(f"\n📝 Chunk {i+1} (Length: {len(chunk.page_content)} characters):\n{chunk.page_content}...\n")

    # Find longest and shortest chunks
    longest_chunk = max(chunks, key=lambda c: len(c.page_content))
    shortest_chunk = min(chunks, key=lambda c: len(c.page_content))

    print(f"\n🔍 Longest Chunk (Length: {len(longest_chunk.page_content)} characters):\n{longest_chunk.page_content[:1500]}...\n")
    print(f"\n🔍 Shortest Chunk (Length: {len(shortest_chunk.page_content)} characters):\n{shortest_chunk.page_content[:1500]}...\n")


def examine_chunks2(chunks):
    """
    Examine the chunks by printing their sizes and checking if overlap is applied.

    Args:
        chunks (list): List of chunked documents.
    """
    print(f"\n📦 Total Chunks: {len(chunks)}")

    for i, chunk in enumerate(chunks[:5]):  # Print only the first 5 chunks for preview
        print(f"\n📝 Chunk {i+1} (Length: {len(chunk.page_content)} characters):")
        print(chunk.page_content[:500])  # Show first 500 chars of chunk

        # If not the first chunk, check overlap
        if i > 0:
            print("\n🔍 Checking Overlap:")
            print("➡ End of previous chunk:", chunks[i - 1].page_content[-150:])  # Last 150 chars
            print("⬅ Start of current chunk:", chunk.page_content[:150])  # First 150 chars

            print("⚠️ Overlap should be visible above (compare both ends)")

    print("\n✅ Chunk examination complete!\n")




def dynamic_chunking(word_docs, min_chunk_size=500, max_chunk_size=1500):
    """
    Dynamically chunks the document based on detected section headers
    while ensuring meaningful groupings and proper chunk sizes.

    Args:
        word_docs (list): Extracted documents from Word.
        min_chunk_size (int): Minimum chunk size before merging with next.
        max_chunk_size (int): Maximum chunk size for a single chunk.

    Returns:
        list: Dynamically chunked documents.
    """
    print("\n🔍 Performing Dynamic Chunking...\n")

    text = word_docs[0].page_content  # Extract full text
    paragraphs = text.split("\n\n")  # Split based on paragraph breaks

    chunks = []
    current_chunk = []

    for i, para in enumerate(paragraphs):
        para = para.strip()

        # **Use same header detection logic from browse_sections**
        is_header = (len(para) < 100 and
                     i < len(paragraphs) - 1 and len(paragraphs[i + 1]) > 100 and
                     (para.isupper() or re.match(r"^[A-Z][a-z].{5,}$", para) or para.endswith(":"))
                    )

        if is_header and current_chunk:
            # If we encounter a new section header and current_chunk has data, finalize the chunk
            chunk_text = "\n\n".join(current_chunk)
            chunks.append(chunk_text)
            current_chunk = []  # Reset for new section

        current_chunk.append(para)

    # Append last chunk if it exists
    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    # Further split large chunks using RecursiveCharacterTextSplitter
    final_chunks = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=max_chunk_size, chunk_overlap=100)

    for chunk in chunks:
        if len(chunk) > max_chunk_size:
            split_chunks = text_splitter.split_text(chunk)
            final_chunks.extend(split_chunks)
        else:
            final_chunks.append(chunk)

    print(f"✅ Total Chunks Created: {len(final_chunks)}")
    
    return final_chunks





def browse_sections(word_docs):
    """
    Dynamically browse through sections of the document by detecting natural section breaks.
    
    Args:
        word_docs (list): Extracted documents from Word.
    
    Returns:
        None: Prints detected sections.
    """
    print("\n🔍 Browsing Document Sections...\n")

    text = word_docs[0].page_content  # Extract full text from LangChain doc
    paragraphs = text.split("\n\n")  # Split into sections based on paragraph breaks

    for i, para in enumerate(paragraphs):
        para = para.strip()

        # **🔹 Rule 1: Section headers are typically short lines**
        if len(para) < 100:  

            # **🔹 Rule 2: Must have a following paragraph (not standalone)**
            if i < len(paragraphs) - 1 and len(paragraphs[i + 1]) > 100:

                # **🔹 Rule 3: Match ALL-CAPS, Title Case, or text ending in ":"**
                if para.isupper() or re.match(r"^[A-Z][a-z].{5,}$", para) or para.endswith(":"):
                    
                    section_title = para.split(":")[0].strip()  # Remove colons from headers
                    print(f"📌 Section Detected: {section_title}\n")

    print("✅ Section Browsing Complete.")



def browse_sectionsOLD(word_docs):
    """
    Browse through sections of the document by detecting headers before chunking.
    
    Args:
        word_docs (list): Extracted documents from Word.
    
    Returns:
        None: Prints detected sections.
    """
    print("\n🔍 Browsing Document Sections...\n")

    text = word_docs[0].page_content  # Extract full text from LangChain doc
    paragraphs = text.split("\n\n")  # Split into sections based on paragraph breaks

    detected_sections = set()  # Track unique sections

    for para in paragraphs:
        para = para.strip()

        # 🔹 Detect Section Headers (Flexible for different formats)
        if re.match(r"^(?=.*\b(Code|Wage Type|Time Type|Attendance|Absence|Policy|Rule|Validation|Exception)\b).{10,}", para) \
           or re.match(r"^[A-Z\s\-()]+:", para):  # Includes all caps headers

            section_title = para.split(":")[0].strip()  # Extract title without `:`
            
            if section_title not in detected_sections:  # Avoid duplicates
                detected_sections.add(section_title)
                print(f"📌 Section Detected: {section_title}\n")
    
    print(f"✅ Total Sections Identified: {len(detected_sections)}")




if __name__ == "__main__":
    """
    Main execution section to test modular functions.
    Comment/uncomment lines to test different parts of the pipeline.
    """

    rootFolder = "/home/rchak007/github/lausdLLM"
    docFolder = os.path.join(rootFolder, "documents")
    file_name = "LLM Att Abs.docx"
    file_path = os.path.join(docFolder, file_name)
    pdf_file_name = "FINAL_Time Type Requirements Document -- FINAL (1).pdf"
    print("File Path:", file_path)

    # Set your document path
    word_file_path = file_path

    # Call the function to read, extract, and convert the document
    word_docs = read_extract_Langdoc(word_file_path)

    # Step 2: Browse Sections Before Chunking
    browse_sections(word_docs)    

    # we wanted to first do Recursive Splitter chunks followed by Token based but looking at how our Document is more like 
    #         requirements style. We might want to try dynamic chunking which chunks at sections instead. So not doing below ones for now.
                    # chunks = chunk_documents(word_docs)  # Step 2: Chunking
                    # # Just for testing examine the chunks.
                    # # examine_chunks(chunks)  # ✅ Examine chunk details
                    # examine_chunks2(chunks)  

                    # # Step 2: Token-based chunking after the RecursiveText splitter.
                    # token_chunks = chunk_TokenTextSplitter(char_chunks)
                    # examine_chunks2(token_chunks) 



    dynamic_chunks = dynamic_chunking(word_docs)
    # # Step 3: Examine chunks (optional)
    examine_chunks(dynamic_chunks)

    # db = create_vector_store(chunks, "/home/rchak007/github/lausdLLM/chromaDocs/chroma/", remove=True)  # Step 3: Embedding
    # qa = initialize_retriever_and_qa(db, chain_type="stuff", k=10)  # Step 4: Retriever & QA

    # You can now pass `word_docs` into your chunking & vector embedding pipeline
    # See readme for chain_type and 
    #     qa is created using LangChain’s ConversationalRetrievalChain, so it returns a chatbot-like response that includes retrieved documents.
    # qa = load_db(word_docs, "stuff", 10, remove=True)  # ✅ Deletes and rebuilds DB
    # qa = load_db(word_docs, "stuff", 10, remove=False)  # ✅ Deletes and rebuilds DB
    # qa = load_db(word_docs, chain_type="map_reduce", k=10, remove=False)  # Uses map_reduce
    # qa = load_db(word_docs, chain_type="refine", k=10, remove=False)  # Uses refine 



    ################# testing 
    # from docx import Document

    # file_path = "/home/rchak007/github/lausdLLM/documents/LLM Att Abs.docx"  # Update path
    # doc = Document(file_path)

    # # Count tables in the document
    # print(f"Total tables found: {len(doc.tables)}")

    # # Print first few rows of the first table (if it exists)
    # if len(doc.tables) > 0:
    #     for row in doc.tables[0].rows[:5]:  # Only first 5 rows
    #         print([cell.text.strip() for cell in row.cells])
