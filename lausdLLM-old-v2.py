
############### START STUFF #########################################################################
#### if we want to remove the database and start over. Esp if i add more documents. and we need to split and embed again
# rm -rf /home/rchak007/github/lausdLLM/chromaDocs/chroma/
#######################################################################################################

import os
import openai
import sys
sys.path.append('../..')
from dotenv import load_dotenv, find_dotenv
from docx import Document

# Load environment variables from .env file
load_dotenv(find_dotenv())

openai.api_key  = os.environ['OPENAI_API_KEY']

print(os.getcwd())  # Prints the current working directory

rootFolder = "/home/rchak007/github/lausdLLM"
docFolder = os.path.join(rootFolder, "documents")
file_name = "LLM Att Abs.docx"
file_path = os.path.join(docFolder, file_name)
pdf_file_name = "FINAL_Time Type Requirements Document -- FINAL (1).pdf"
print("File Path:", file_path)



from langchain.document_loaders import PyPDFLoader
# Combine the root folder and file name
pdf_file_path = os.path.join(docFolder, pdf_file_name)
print("File Path:", pdf_file_path)
# Load the PDF
loader = PyPDFLoader(pdf_file_path)

pages = loader.load()

len(pages)

page = pages[5]

print(page.page_content[0:500])

page.metadata

from langchain_community.document_loaders.blob_loaders.file_system import FileSystemBlobLoader
from langchain.document_loaders.generic import GenericLoader
from langchain.document_loaders.parsers import OpenAIWhisperParser
from langchain.document_loaders.blob_loaders.youtube_audio import YoutubeAudioLoader

from langchain.document_loaders import WebBaseLoader



from langchain.document_loaders import UnstructuredWordDocumentLoader

# Then download NLTK resources
import nltk
nltk.download('punkt')
nltk.download('punkt_tab')
nltk.download('averaged_perceptron_tagger')
nltk.download('averaged_perceptron_tagger_eng')

def extract_text_from_word(file_path):
    doc = Document(file_path)
    extracted_text = []

    for para in doc.paragraphs:
        # Identify bullet points and convert them to readable text
        if para.style.name.startswith("List"):
            extracted_text.append(f"- {para.text}")  # Format bullet points
        else:
            extracted_text.append(para.text)

    # Extract tables
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(row_text))  # Format table as text
        extracted_text.append("\n".join(table_text))  # Add table text to document

    return "\n".join(extracted_text)  # Convert list to a single string


# Path to your Word document
word_file_path = "/home/rchak007/github/lausdLLM/documents/LLM Att Abs.docx"
# "/content/drive/MyDrive/Colab Notebooks/AI-LAUSD/Documents/Attendance and Absences Requirements.docx"
# word_file_path1 = "/content/drive/MyDrive/Colab Notebooks/AI-LAUSD/Documents/Time Type Requirements Document -- FINAL.docx"
# word_file_path2 = "/content/drive/MyDrive/Colab Notebooks/AI-LAUSD/Documents/AB2160 Requirements Document 1GXX Classified.doc"
# word_file_path3 = "/content/drive/MyDrive/Colab Notebooks/AI-LAUSD/Documents/Absence Quotas Requirements Document.docx"

# # Load the Word document
# loader = UnstructuredWordDocumentLoader(word_file_path)
# loader = UnstructuredWordDocumentLoader(word_file_path1)
# loader = UnstructuredWordDocumentLoader(word_file_path2)
# loader = UnstructuredWordDocumentLoader(word_file_path3)
# word_document = loader.load()

# Extract structured text from Word
word_text = extract_text_from_word(word_file_path)

# Convert extracted text into a LangChain Document
from langchain.schema import Document as LC_Document
wordDocs = [LC_Document(page_content=word_text, metadata={"source": word_file_path})]

# # Create separate loaders for each document
# loader1 = UnstructuredWordDocumentLoader(word_file_path)
# # loader2 = UnstructuredWordDocumentLoader(word_file_path1)
# # loader3 = UnstructuredWordDocumentLoader(word_file_path2)
# # loader4 = UnstructuredWordDocumentLoader(word_file_path3)

# # Load all documents and combine their contents
# wordDocs = []
# wordDocs.extend(loader1.load())
# wordDocs.extend(loader2.load())
# wordDocs.extend(loader3.load())
# wordDocs.extend(loader4.load())

# Now 'docs' contains all documents
print(f"Total number of documents loaded: {len(wordDocs)}")

len(wordDocs)

print(wordDocs[0].page_content[0:200])


# # Document Splitting
from langchain.text_splitter import RecursiveCharacterTextSplitter, CharacterTextSplitter
# #### CharacterTextSplitter in LangChain blindly chunks text based on the number of characters, without considering sentence structure or semantic meaning. thats why we will not use this

text_splitter_ch = CharacterTextSplitter(
    separator="\n",
    chunk_size=1000,
    chunk_overlap=150,
    length_function=len
)


# ### For smarter splitting, RecursiveCharacterTextSplitter is better. It tries to chunk at meaningful breakpoints, like paragraphs or sentences, before falling back on character-level splitting.
# ### below claude recommendation after warning

# Define the first splitter to chunk meaningfully
recursive_splitter  = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=150,
    length_function=len,
    separators=["\n\n", "\n", " ", ""]  # Will try these separators in order
)

first_chunks  = recursive_splitter.split_documents(wordDocs)
print(f"After Recursive Split: {len(first_chunks)} chunks")

len(first_chunks )

len(wordDocs)


# # Token splitting
# ### The TokenTextSplitter in LangChain chunks text based on tokens rather than raw characters. It ensures that splits align better with natural language structures compared to CharacterTextSplitter, which blindly splits by character count.
# ### Uses a Tokenizer
# 
# - It splits text based on tokens rather than characters.
# - The tokenizer depends on the underlying model (e.g., OpenAI's tokenizer for GPT).

# In[41]:


from langchain.text_splitter import TokenTextSplitter


# A chunk_size of 10 is far too small for this document because it’s token-based, and tokens are typically smaller than words. Given the nature of your text (policy-based, structured, with some formulae), an ideal chunk_size should be large enough to capture full logical sections while avoiding excessive length.
# 
# Recommended Chunk Size for Your Document
# 100–200 tokens: Captures key ideas while keeping chunks digestible.
# 200–300 tokens: Ensures full sentences and structured sections stay intact.
# 300–400 tokens: Best if you want to keep entire policy subsections in one chunk.
# Since your document has structured bullet points, formulas, and policy descriptions, I'd suggest starting with chunk_size=256 and chunk_overlap=50 to retain key information across chunks.



token_splitter = TokenTextSplitter(chunk_size=1000, chunk_overlap=200)
final_chunks  = token_splitter.split_documents(first_chunks)

final_chunks [0]
final_chunks [0].metadata


# ## Context aware splitting - not doing this as it was for Markdown

# # Embeddings
# Let's take our final_chunks  and embed them.
# 
# 
# Your Embedding Model May Be Weak or Mismatched
# If the embedding model does not properly capture the semantic meaning of your documents, the similarity search might return irrelevant or overly short results.
# 
# Solution: Try using a better embedding model (e.g., OpenAI’s text-embedding-ada-002 or a larger SentenceTransformers model like all-MiniLM-L6-v2).

from langchain.embeddings.openai import OpenAIEmbeddings
# embedding = OpenAIEmbeddings()
embedding = OpenAIEmbeddings(model="text-embedding-ada-002")
# embedding = OpenAIEmbeddings(model="all-MiniLM-L6-v2")

import numpy as np

from langchain.vectorstores import Chroma

# get_ipython if for Jupiter notebook only
# get_ipython().system('rm -rf /home/rchak007/github/lausdLLM/chromaDocs/chroma  # remove old database')
import os

# try:
#     get_ipython().system('rm -rf /home/rchak007/github/lausdLLM/chromaDocs/chroma')
# except NameError:
#     import os
#     os.system('rm -rf /home/rchak007/github/lausdLLM/chromaDocs/chroma')

# In[50]:


#### if we want to remove the database and start over. Esp if i add more documents.
# rm -rf /home/rchak007/github/lausdLLM/chromaDocs/chroma/

# Create the path for ChromaDB storage
persist_directory = '/home/rchak007/github/lausdLLM/chromaDocs/chroma/'
os.makedirs(persist_directory, exist_ok=True)  # Ensure fresh storage

# It's good practice to create the directory if it doesn't exist
import os
os.makedirs(persist_directory, exist_ok=True)

# Check if the database already exists
if os.path.exists(persist_directory) and os.listdir(persist_directory):
    print("Loading existing vector database...")
    vectordb = Chroma(persist_directory=persist_directory, embedding_function=OpenAIEmbeddings())
else:
    print("Creating new vector database...")
    vectordb = Chroma.from_documents(documents=final_chunks , embedding=OpenAIEmbeddings(), persist_directory=persist_directory)
    
vectordb.persist()  # Ensure changes are saved


print(vectordb._collection.count())

question = "give me some info on vacation "

docs = vectordb.similarity_search(question,k=10)


len(docs)


docs[0].page_content
# # Retrieval
# 
# Retrieval is the centerpiece of our retrieval augmented generation (RAG) flow.
# 
# Let's get our vectorDB from before.
# 
# ## When a query comes in you want to retrieve the most relevant splits.
print(vectordb._collection.count())
question = "give me some info on vacation quota"

# ## Addressing Diversity: Maximum marginal relevance
# Last class we introduced one problem: how to enforce diversity in the search results.
# 
# `Maximum marginal relevance` strives to achieve both relevance to the query *and diversity* among the results.

# docs_ss = vectordb.similarity_search(question,k=3)
docs_ss = vectordb.similarity_search(question,k=10)
docs_ss[0].page_content[:1200]
docs_ss[1].page_content[:100]
docs_ss[2].page_content[:100]
docs_ss[3].page_content[:100]
docs_ss[4].page_content[:100]
docs_mmr = vectordb.max_marginal_relevance_search(question,k=10)
docs_mmr[0].page_content[:100]
docs_mmr[1].page_content[:100]
docs_mmr[2].page_content[:100]
docs_mmr[3].page_content[:100]
docs_mmr[4].page_content[:100]
docs_mmr[5].page_content[:100]
docs_mmr[6].page_content[:100]

# ## Addressing Specificity: working with metadata
# In last lecture, we showed that a question about the third lecture can include results from other lectures as well.
# 
# To address this, many vectorstores support operations on `metadata`.
# 
# `metadata` provides context for each embedded chunk.
question
docs = vectordb.similarity_search(
    question,
    k=3,
    # filter={"source":"docs/cs229_lectures/MachineLearning-Lecture03.pdf"}
    filter={"source":"/home/rchak007/github/lausdLLM/documents/LLM Att Abs.docx"}
)
for d in docs:
    print(d.metadata)
# ## Addressing Specificity: working with metadata using self-query retriever
# But we have an interesting challenge: we often want to infer the metadata from the query itself.
# 
# To address this, we can use `SelfQueryRetriever`, which uses an LLM to extract:
# 
# 1. The `query` string to use for vector search
# 2. A metadata filter to pass in as well
# 
# Most vector databases support metadata filters, so this doesn't require any new databases or indexes.

# In[80]:


from langchain.llms import OpenAI
from langchain.retrievers.self_query.base import SelfQueryRetriever
from langchain.chains.query_constructor.base import AttributeInfo


# In[81]:


metadata_field_info = [
    AttributeInfo(
        name="source",
        description="The lecture the chunk is from, should be one of `/home/rchak007/github/lausdLLM/documents/LLM Att Abs.docx`",
        type="string",
    ),
    AttributeInfo(
        name="page",
        description="The page from the document",
        type="integer",
    ),
]


# Note: The default model for OpenAI ("from langchain.llms import OpenAI") is text-davinci-003. Due to the deprication of OpenAI's model text-davinci-003 on 4 January 2024, you'll be using OpenAI's recommended replacement model gpt-3.5-turbo-instruct instead.

# In[82]:


document_content_description = "Quota info"
llm = OpenAI(model='gpt-3.5-turbo-instruct', temperature=0)
retriever = SelfQueryRetriever.from_llm(
    llm,
    vectordb,
    document_content_description,
    metadata_field_info,
    verbose=True
)


# You will receive a warning about predict_and_parse being deprecated the first time you executing the next line. This can be safely ignored.

# In[83]:


docs = retriever.get_relevant_documents(question)
# docs = retriever.invoke(question)


# In[84]:


question = "describe PROFESSIONAL DEVELOPMENT TIME" 


# In[85]:


for d in docs:
    print(d.metadata)


# ## Additional tricks: compression
# Another approach for improving the quality of retrieved docs is compression.
# 
# Information most relevant to a query may be buried in a document with a lot of irrelevant text.
# 
# Passing that full document through your application can lead to more expensive LLM calls and poorer responses.
# 
# Contextual compression is meant to fix this.

# In[86]:


from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor


# In[87]:


def pretty_print_docs(docs):
    print(f"\n{'-' * 100}\n".join([f"Document {i+1}:\n\n" + d.page_content for i, d in enumerate(docs)]))


# In[88]:


# Wrap our vectorstore
llm = OpenAI(temperature=0, model="gpt-3.5-turbo-instruct")
compressor = LLMChainExtractor.from_llm(llm)


# In[89]:


compression_retriever = ContextualCompressionRetriever(
    base_compressor=compressor,
    base_retriever=vectordb.as_retriever()
)


# In[90]:


question = "what did they say about MISCELLANEOUS Time?"
compressed_docs = compression_retriever.get_relevant_documents(question)
pretty_print_docs(compressed_docs)


# ### Combining various techniques

# In[91]:


compression_retriever = ContextualCompressionRetriever(
    base_compressor=compressor,
    base_retriever=vectordb.as_retriever(search_type = "mmr")
)


# In[92]:


compressed_docs = compression_retriever.get_relevant_documents(question)
pretty_print_docs(compressed_docs)


# ## Other types of retrieval
# It's worth noting that vectordb as not the only kind of tool to retrieve documents.
# 
# The LangChain retriever abstraction includes other ways to retrieve documents, such as TF-IDF or SVM.

# In[93]:


from langchain.retrievers import SVMRetriever
from langchain.retrievers import TFIDFRetriever
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter


# In[94]:


len(splits)


# In[95]:


all_page_text=[p.page_content for p in wordDocs]
joined_page_text=" ".join(all_page_text)


# In[96]:


# Split
other_text_splitter = RecursiveCharacterTextSplitter(chunk_size = 1500,chunk_overlap = 150)
other_splits = text_splitter.split_text(joined_page_text)


# In[97]:


# Retrieve
svm_retriever = SVMRetriever.from_texts(other_splits,embedding)
tfidf_retriever = TFIDFRetriever.from_texts(other_splits)


# In[98]:


question = "describe PROFESSIONAL DEVELOPMENT attendance?"
docs_svm=svm_retriever.get_relevant_documents(question)
docs_svm[0]


# In[99]:


docs_tfidf=tfidf_retriever.get_relevant_documents(question)
docs_tfidf[0]


# ## Question Answering

# In[100]:


import datetime
current_date = datetime.datetime.now().date()
if current_date < datetime.date(2023, 9, 2):
    llm_name = "gpt-3.5-turbo-0301"
else:
    llm_name = "gpt-3.5-turbo"
print(llm_name)


# In[101]:


print(vectordb._collection.count())


# In[102]:


question = "what did they say about MISCELLANEOUS Time?"
docs = vectordb.similarity_search(question,k=3)
len(docs)


# In[103]:


from langchain.chat_models import ChatOpenAI
llm = ChatOpenAI(model_name=llm_name, temperature=0)


# ### RetrievalQA chain

# In[104]:


from langchain.chains import RetrievalQA


# In[105]:


qa_chain = RetrievalQA.from_chain_type(
    llm,
    retriever=vectordb.as_retriever()
)


# In[106]:


result = qa_chain({"query": question})


# In[107]:


result["result"]


# ### Prompt

# In[108]:


from langchain.prompts import PromptTemplate

# Build prompt
template = """Use the following pieces of context to answer the question at the end. If you don't know the answer, just say that you don't know, don't try to make up an answer. Use three sentences maximum. Keep the answer as concise as possible. Always say "thanks for asking!" at the end of the answer. 
{context}
Question: {question}
Helpful Answer:"""
QA_CHAIN_PROMPT = PromptTemplate.from_template(template)


# In[109]:


# Run chain
qa_chain = RetrievalQA.from_chain_type(
    llm,
    retriever=vectordb.as_retriever(),
    return_source_documents=True,
    chain_type_kwargs={"prompt": QA_CHAIN_PROMPT}
)


# In[110]:


result = qa_chain({"query": question})


# In[111]:


result["result"]


# In[112]:


result["source_documents"][0]


# ### RetrievalQA chain types

# In[113]:


qa_chain_mr = RetrievalQA.from_chain_type(
    llm,
    retriever=vectordb.as_retriever(),
    chain_type="map_reduce"
)


# In[114]:


result = qa_chain_mr({"query": question})


# In[115]:


result["result"]


# ## TO DO 
# If you wish to experiment on the LangSmith platform (previously known as LangChain Plus):
# 
# Go to LangSmith and sign up
# Create an API key from your account's settings
# Use this API key in the code below
# uncomment the code
# Note, the endpoint in the video differs from the one below. Use the one below.

# ## TO DO 
# ### RetrievalQA limitations
#  
# QA fails to preserve conversational history.

# # 05-Chat

# In[116]:


from langchain.vectorstores import Chroma
from langchain.embeddings.openai import OpenAIEmbeddings
# persist_directory = 'docs/chroma/'
# embedding = OpenAIEmbeddings()
# vectordb = Chroma(persist_directory=persist_directory, embedding_function=embedding)

# already loaded all this.


# In[117]:


question


# In[118]:


# question = "what did they say about MISCELLANEOUS Time?"
docs = vectordb.similarity_search(question,k=3)
len(docs)


# In[119]:


from langchain.chat_models import ChatOpenAI
llm = ChatOpenAI(model_name=llm_name, temperature=0)
llm.predict("Hello world!")


# In[120]:


# Build prompt
from langchain.prompts import PromptTemplate
template = """Use the following pieces of context to answer the question at the end. If you don't know the answer, just say that you don't know, don't try to make up an answer. Use three sentences maximum. Keep the answer as concise as possible. Always say "thanks for asking!" at the end of the answer. 
{context}
Question: {question}
Helpful Answer:"""
QA_CHAIN_PROMPT = PromptTemplate(input_variables=["context", "question"],template=template,)

# Run chain
from langchain.chains import RetrievalQA
# question = "Is probability a class topic?"
question = "what did they say about MISCELLANEOUS Time?"
qa_chain = RetrievalQA.from_chain_type(llm,
                                       retriever=vectordb.as_retriever(),
                                       return_source_documents=True,
                                       chain_type_kwargs={"prompt": QA_CHAIN_PROMPT})


result = qa_chain({"query": question})
result["result"]


# ### Memory

# In[121]:


from langchain.memory import ConversationBufferMemory
memory = ConversationBufferMemory(
    memory_key="chat_history",
    return_messages=True
)


# ### ConversationalRetrievalChain

# In[122]:


from langchain.chains import ConversationalRetrievalChain
retriever=vectordb.as_retriever()
qa = ConversationalRetrievalChain.from_llm(
    llm,
    retriever=retriever,
    memory=memory
)


# In[123]:


result = qa({"question": question})


# In[124]:


result['answer']


# In[125]:


question = "can you expand on the rules for annual physical examinations"
result = qa({"question": question})


# In[126]:


result['answer']


# ## Create a chatbot that works on your documents

# In[127]:


# os.environ["LANG_SMITH_API"]


# In[128]:


import os
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_ENDPOINT"] = "https://api.langchain.plus"
LANG_SMITH_API = os.environ["LANG_SMITH_API"]
# openai.api_key  = os.environ['OPENAI_API_KEY']


# In[129]:


from langchain_openai import ChatOpenAI

llm = ChatOpenAI()
llm.invoke("Hello, world!")


# In[130]:


from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter, RecursiveCharacterTextSplitter
from langchain.vectorstores import DocArrayInMemorySearch
from langchain.document_loaders import TextLoader
from langchain.chains import RetrievalQA,  ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import TextLoader
from langchain.document_loaders import PyPDFLoader


# In[139]:


def load_db(file, chain_type, k):


    persist_directory = "/home/rchak007/github/lausdLLM/chromaDocs/chroma/"  # Persistent storage for embeddings
    
    # Check if Chroma DB exists
    if os.path.exists(persist_directory) and os.listdir(persist_directory):
        print("Loading existing Chroma vector database...")
        db = Chroma(persist_directory=persist_directory, embedding_function=OpenAIEmbeddings())
    else:
        print("Creating new Chroma vector database...")
        # Load documents
        loader = UnstructuredWordDocumentLoader(file)
        documents = loader.load()

        # Split documents
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
        docs = text_splitter.split_documents(documents)

        # Create vector store
        embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
        db = Chroma.from_documents(docs, embeddings, persist_directory=persist_directory)
        db.persist()  # Save for future use


    retriever = db.as_retriever(search_type="mmr", search_kwargs={"k": k, "fetch_k": 20})
    # create a chatbot chain. Memory is managed externally.
    qa = ConversationalRetrievalChain.from_llm(
        llm=ChatOpenAI(model_name=llm_name, temperature=0), 
        chain_type=chain_type, 
        retriever=retriever, 
        return_source_documents=True,
        return_generated_question=True,
    )
    return qa 


# In[168]:


import panel as pn
import param

class cbfs(param.Parameterized):
    chat_history = param.List([])
    answer = param.String("")
    db_query  = param.String("")
    db_response = param.List([])
    
    def __init__(self,  **params):
        super(cbfs, self).__init__( **params)
        self.panels = []
        # self.loaded_file = "docs/cs229_lectures/MachineLearning-Lecture01.pdf"
        self.loaded_file = "/home/rchak007/github/lausdLLM/documents/LLM Att Abs.docx"
        self.qa = load_db(self.loaded_file,"stuff", 10)
    
    def call_load_db(self, count):
        if count == 0 or file_input.value is None:  # init or no file specified :
            return pn.pane.Markdown(f"Loaded File: {self.loaded_file}")
        else:
            file_input.save("temp.pdf")  # local copy
            self.loaded_file = file_input.filename
            button_load.button_style="outline"
            self.qa = load_db("temp.pdf", "stuff", 10)
            button_load.button_style="solid"
        self.clr_history()
        return pn.pane.Markdown(f"Loaded File: {self.loaded_file}")


    def convchain(self, query):
        if not query:
            return pn.WidgetBox(pn.Row('User:', pn.pane.Markdown("", width=600)), scroll=True)
        
        try:
            result = self.qa({"question": query, "chat_history": self.chat_history})
    
            # Check if the result contains an answer
            if "answer" not in result or not result["answer"]:
                self.answer = "Sorry, I couldn't find an answer to that question."
            else:
                self.answer = result["answer"]

            # Log and display the retrieved chunks
            retrieved_chunks = result.get("source_documents", [])
            num_chunks = len(retrieved_chunks)
            chunk_texts = [f"Chunk {i+1}: {doc.page_content[:300]}..." for i, doc in enumerate(retrieved_chunks)]     
    
            self.db_query = query
            self.db_response = chunk_texts  # Store retrieved chunks for later debugging    

            # Update chat history
            self.chat_history.append((query, self.answer))
    
            # Update UI
            # self.panels.extend([
            #     pn.Row('User:', pn.pane.Markdown(query, width=600)),
            #     pn.Row('ChatBot:', pn.pane.Markdown(self.answer, width=600, background='#F6F6F6'))
            # ])
            # self.panels.extend([
            #     pn.Row('User:', pn.pane.Markdown(query, width=600)),
            #     pn.Row('ChatBot:', pn.pane.Markdown(self.answer, width=600))  # No background or style
            # ])
            # Update UI
            self.panels.extend([
                pn.Row('User:', pn.pane.Markdown(query, width=600)),
                pn.Row('ChatBot:', pn.pane.Markdown(self.answer, width=600)),
                pn.Row('Retrieved Chunks ********************  :', pn.pane.Markdown(f"Total Chunks Retrieved: {num_chunks}", width=600)),
                pn.Column(*[pn.pane.Markdown(chunk, width=600) for chunk in chunk_texts])  # Display chunk previews
            ])


        except Exception as e:
            # Handle unexpected errors and show them in UI
            error_msg = f"An error occurred: {str(e)}"
            self.panels.append(pn.Row('ChatBot:', pn.pane.Markdown(error_msg, width=600, background='#FFCCCC')))
    
        return pn.WidgetBox(*self.panels, scroll=True)
    
#     def convchain(self, query):
#         if not query:
#             return pn.WidgetBox(pn.Row('User:', pn.pane.Markdown("", width=600)), scroll=True)
#         result = self.qa({"question": query, "chat_history": self.chat_history})

#         # Debugging outputs
#         print(f"Query: {query}")
#         print(f"Result from QA: {result}")
        
#         self.chat_history.extend([(query, result["answer"])])
#         self.db_query = result["generated_question"]
#         self.db_response = result["source_documents"]
#         self.answer = result['answer'] 
#         # More debugging to confirm the answer is set
#         print(f"ChatBot Answer: {self.answer}")        
#         # self.panels.extend([
#         #     pn.Row('User:', pn.pane.Markdown(query, width=600)),
#         #     pn.Row('ChatBot:', pn.pane.Markdown(self.answer, width=600, style={'background-color': '#F6F6F6'}))
#         # ])
#         self.panels.extend([
#         pn.Row('User:', pn.pane.Markdown(query, width=600)),
#         pn.Row('ChatBot:', pn.pane.Markdown(self.answer, width=600))  # Removed the 'style' argument
# ])

#         inp.value = ''  #clears loading indicator when cleared
#         return pn.WidgetBox(*self.panels,scroll=True)

    @param.depends('db_query ', )
    def get_lquest(self):
        if not self.db_query :
            return pn.Column(
                pn.Row(pn.pane.Markdown(f"Last question to DB:", styles={'background-color': '#F6F6F6'})),
                pn.Row(pn.pane.Str("no DB accesses so far"))
            )
        return pn.Column(
            pn.Row(pn.pane.Markdown(f"DB query:", styles={'background-color': '#F6F6F6'})),
            pn.pane.Str(self.db_query )
        )

    @param.depends('db_response', )
    def get_sources(self):
        if not self.db_response:
            return pn.WidgetBox(pn.Row(pn.pane.Markdown(f"Last Sources from DB lookup:", styles={'background-color': '#F6F6F6'})),
                            pn.Row(pn.pane.Str("No chunks retrieved")))
        rlist=[pn.Row(pn.pane.Markdown(f"Result of DB lookup:", styles={'background-color': '#F6F6F6'}))]
        for doc in self.db_response:
            rlist.append(pn.Row(pn.pane.Str(doc)))
        return pn.WidgetBox(*rlist, width=600, scroll=True)

    @param.depends('convchain', 'clr_history') 
    def get_chats(self):
        if not self.chat_history:
            return pn.WidgetBox(pn.Row(pn.pane.Str("No History Yet")), width=600, scroll=True)
        rlist=[pn.Row(pn.pane.Markdown(f"Current Chat History variable", styles={'background-color': '#F6F6F6'}))]
        for exchange in self.chat_history:
            rlist.append(pn.Row(pn.pane.Str(exchange)))
        return pn.WidgetBox(*rlist, width=600, scroll=True)

    def clr_history(self,count=0):
        self.chat_history = []
        return 


# In[170]:


cb = cbfs()

file_input = pn.widgets.FileInput(accept='.docx')
button_load = pn.widgets.Button(name="Load DB", button_type='primary')
button_clearhistory = pn.widgets.Button(name="Clear History", button_type='warning')
button_clearhistory.on_click(cb.clr_history)
inp = pn.widgets.TextInput( placeholder='Enter text here…')

bound_button_load = pn.bind(cb.call_load_db, button_load.param.clicks)
conversation = pn.bind(cb.convchain, inp) 

# jpg_pane = pn.pane.Image( './img/convchain.jpg')


conversation_panel = pn.Column(cb.convchain(''))  # Initialize with an empty conversation


# Create the Send Question button
send_button = pn.widgets.Button(name="Send Question", button_type='primary')
# Container to hold conversation history (for UI display)
# conversation_panel = pn.Column()  # This will dynamically update with chat history
# conversation_panel = pn.bind(cb.convchain, inp)

def on_send_click(event):
    conversation_panel.objects = [cb.convchain(inp.value)]
    # conversation_panel.append(cb.convchain(inp.value))
    inp.value = ''  # Clear input after sending



# Link the button click to the function
send_button.on_click(on_send_click)

# Function to dynamically display chat history in the UI
def display_conversation():
    conversation_panel.clear()  # Clear previous entries
    for user_query, bot_response in cb.chat_history:
        conversation_panel.append(pn.Row('User:', pn.pane.Markdown(user_query, width=600)))
        conversation_panel.append(pn.Row('ChatBot:', pn.pane.Markdown(bot_response, width=600, style={'background-color': '#F6F6F6'})))

# Bind the display function to update whenever chat history changes
pn.bind(display_conversation)



# Widget to display the chat output
chat_output = pn.pane.Markdown("", width=600)


def clear_history(event):
    cb.clr_history()
    conversation_panel.objects = [cb.convchain('')]  # Reset the conversation panel



button_clearhistory.on_click(clear_history)



tab1 = pn.Column(
    pn.Row(inp, send_button),  # Input field with Send button
    pn.layout.Divider(),
    # pn.panel(cb.convchain, parameters=['chat_history'], height=300),  # Display the conversation dynamically
    conversation_panel,
    pn.layout.Divider(),
)

# # Update the layout for Conversation tab
# tab1 = pn.Column(
#     pn.Row(inp, send_button),  # Input field with Send button
#     pn.layout.Divider(),
#     conversation_panel,
#     pn.layout.Divider(),
# )


tab2= pn.Column(
    pn.panel(cb.get_lquest),
    pn.layout.Divider(),
    pn.panel(cb.get_sources ),
)
tab3= pn.Column(
    pn.panel(cb.get_chats),
    pn.layout.Divider(),
)
tab4=pn.Column(
    pn.Row( file_input, button_load, bound_button_load),
    pn.Row( button_clearhistory, pn.pane.Markdown("Clears chat history. Can use to start a new topic" )),
    pn.layout.Divider(),
    # pn.Row(jpg_pane.clone(width=400))
)
dashboard = pn.Column(
    pn.Row(pn.pane.Markdown('# ChatWithYourData_Bot')),
    pn.Tabs(('Conversation', tab1), ('Database', tab2), ('Chat History', tab3),('Configure', tab4))
)
# dashboard


# In[158]:


# import panel as pn
# pn.extension()


# In[159]:


# dashboard


# In[ ]:



if __name__.startswith("bokeh"):
    dashboard.servable()

